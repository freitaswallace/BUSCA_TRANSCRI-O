#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Sistema de Busca Avançada em Arquivos Word com IA
Utiliza Google Gemini para identificação inteligente de nomes e empresas
Autor: Gerado automaticamente
Versão: 2.0
"""

import os
import sys
import json
import time
import threading
import queue
import unicodedata
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple
import tkinter as tk
from tkinter import messagebox
import customtkinter as ctk

# Bibliotecas para manipulação de Word
try:
    from docx import Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
except ImportError:
    print("ERRO: python-docx não instalado. Execute: pip install python-docx")
    sys.exit(1)

# Biblioteca para ler arquivos .doc antigos (formato binário) - Windows COM
try:
    import win32com.client
    import pythoncom
    HAS_WORD_COM = True
except ImportError:
    HAS_WORD_COM = False
    print("AVISO: pywin32 não instalado. Arquivos .doc antigos podem não funcionar.")
    print("Execute: pip install pywin32")

# Biblioteca Google Gemini
try:
    import google.generativeai as genai
except ImportError:
    print("ERRO: google-generativeai não instalado. Execute: pip install google-generativeai")
    sys.exit(1)


# ===========================
# CONFIGURAÇÕES GLOBAIS
# ===========================
PASTA_BASE = r"\\192.168.20.100\trabalho\Transcrições"
CONFIG_FILE = "config.json"
EXTENSIONS = ['.docx', '.doc']
NUM_THREADS = 10

# Cores (Tema Claro Agradável)
COLORS = {
    "bg_main": "#F5F7FA",           # Fundo principal (azul acinzentado muito claro)
    "bg_card": "#FFFFFF",            # Cards e frames (branco puro)
    "bg_header": "#4A90E2",          # Cabeçalho (azul suave)
    "bg_input": "#ECF0F1",           # Inputs (cinza muito claro)
    "fg_primary": "#2C3E50",         # Texto primário (cinza escuro azulado)
    "fg_secondary": "#7F8C8D",       # Texto secundário (cinza médio)
    "fg_header": "#FFFFFF",          # Texto do cabeçalho (branco)
    "accent": "#5DADE2",             # Botão principal (azul claro)
    "accent_hover": "#3498DB",       # Hover do botão (azul médio)
    "success": "#27AE60",            # Verde suave
    "error": "#E74C3C",              # Vermelho suave
    "warning": "#F39C12",            # Laranja suave
    "info": "#3498DB",               # Azul informação
    "border": "#BDC3C7"              # Bordas sutis
}


# ===========================
# FUNÇÕES AUXILIARES
# ===========================
def normalize_text(text: str) -> str:
    """
    Remove acentos e normaliza texto para busca
    ANTONIO CARDINÁ -> ANTONIO CARDINA
    """
    if not text:
        return ""

    # Normaliza para NFD (decompõe caracteres acentuados)
    nfd = unicodedata.normalize('NFD', text)
    # Remove marcas diacríticas (acentos)
    without_accents = ''.join(char for char in nfd if unicodedata.category(char) != 'Mn')
    # Converte para maiúsculas para comparação case-insensitive
    result = without_accents.upper()

    # Remove espaços extras
    result = ' '.join(result.split())

    return result


def extract_text_from_old_doc(file_path: str) -> str:
    """
    Extrai texto de arquivos .doc antigos (formato binário)
    Usa Microsoft Word via COM Automation (Windows)
    """
    if not HAS_WORD_COM:
        print(f"[DEBUG] pywin32 não disponível, não é possível ler .doc antigo")
        return ""

    # CRÍTICO: Inicializar COM para esta thread
    pythoncom.CoInitialize()

    word_app = None
    doc = None
    try:
        print(f"[DEBUG] Abrindo .doc com Word COM: {file_path}")

        # Criar instância do Word
        word_app = win32com.client.Dispatch("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = 0  # Não mostrar alertas

        # Abrir documento (usar caminho absoluto)
        abs_path = os.path.abspath(file_path)
        doc = word_app.Documents.Open(abs_path, ReadOnly=True)

        # Extrair todo o texto
        full_text = doc.Content.Text

        # Fechar documento
        doc.Close(False)
        word_app.Quit()

        print(f"[DEBUG] Texto extraído com sucesso ({len(full_text)} caracteres)")
        return full_text

    except Exception as e:
        print(f"[DEBUG] Erro ao extrair texto de .doc antigo com Word COM: {e}")

        # Tentar fechar Word se ainda estiver aberto
        try:
            if doc:
                doc.Close(False)
            if word_app:
                word_app.Quit()
        except:
            pass

        return ""

    finally:
        # CRÍTICO: Sempre desinicializar COM
        pythoncom.CoUninitialize()


def is_old_doc_format(file_path: str) -> bool:
    """
    Verifica se o arquivo é .doc antigo (formato binário)
    """
    try:
        # Tentar abrir com python-docx
        doc = Document(file_path)
        return False  # É .docx (formato novo)
    except Exception as e:
        # Se der erro, provavelmente é .doc antigo
        error_msg = str(e).lower()
        if 'not a word file' in error_msg or 'content type' in error_msg:
            return True
        return False


# ===========================
# CLASSE DE CONFIGURAÇÃO
# ===========================
class ConfigManager:
    """Gerencia a persistência da API Key"""

    def __init__(self, config_file: str = CONFIG_FILE):
        self.config_file = config_file
        self.config = self.load_config()

    def load_config(self) -> Dict:
        """Carrega configurações do arquivo JSON"""
        if os.path.exists(self.config_file):
            try:
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                print(f"Erro ao carregar config: {e}")
                return {}
        return {}

    def save_config(self, config: Dict) -> bool:
        """Salva configurações no arquivo JSON"""
        try:
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(config, f, indent=4, ensure_ascii=False)
            return True
        except Exception as e:
            print(f"Erro ao salvar config: {e}")
            return False

    def get_api_key(self) -> str:
        """Retorna a API Key salva"""
        return self.config.get('gemini_api_key', '')

    def set_api_key(self, api_key: str) -> bool:
        """Define e salva a API Key"""
        self.config['gemini_api_key'] = api_key
        return self.save_config(self.config)


# ===========================
# CLASSE DO MOTOR DE BUSCA
# ===========================
class WordSearchEngine:
    """Motor de busca paralelo para arquivos Word"""

    def __init__(self, base_path: str, num_threads: int = NUM_THREADS):
        self.base_path = base_path
        self.num_threads = num_threads
        self.results_queue = queue.Queue()
        self.progress_queue = queue.Queue()
        self.stop_event = threading.Event()
        self.files_found = []
        self.files_with_errors = []
        self.total_files_processed = 0
        self.lock = threading.Lock()
        # Semáforo para limitar chamadas simultâneas à IA (máximo 3 por vez)
        self.ai_semaphore = threading.Semaphore(3)

    def search_in_document(self, file_path: str, search_term: str, use_ai: bool = False,
                          api_key: str = None) -> Tuple[bool, str]:
        """
        Busca um termo em um documento Word
        Prioriza termos em negrito e sublinhado
        Ignora acentos na comparação
        Busca flexível: encontra mesmo com palavras no meio
        Suporta .doc (antigo) e .docx (novo)
        Retorna (encontrado: bool, contexto: str)
        """
        try:
            # Verificar se é .doc antigo (formato binário)
            if is_old_doc_format(file_path):
                print(f"[DEBUG] Arquivo .doc ANTIGO detectado! Usando extração de texto...")
                return self.search_in_old_doc(file_path, search_term, use_ai, api_key)

            # Arquivo .docx moderno
            doc = Document(file_path)
            # Normalizar termo de busca (remove acentos e converte para maiúsculas)
            search_term_normalized = normalize_text(search_term)

            # Dividir termo em palavras para busca flexível
            search_words = search_term_normalized.split()

            found_contexts = []

            print(f"\n[DEBUG] ==========================================")
            print(f"[DEBUG] Buscando '{search_term}'")
            print(f"[DEBUG] Normalizado: '{search_term_normalized}'")
            print(f"[DEBUG] Palavras: {search_words}")
            print(f"[DEBUG] Total de parágrafos: {len(doc.paragraphs)}")
            print(f"[DEBUG] ==========================================\n")

            # Mostrar primeiros parágrafos para debug
            sample_paras = []
            for i, p in enumerate(doc.paragraphs[:5]):
                if p.text.strip():
                    sample_paras.append(f"  [{i}] {p.text[:80]}")
            if sample_paras:
                print(f"[DEBUG] Primeiros parágrafos do documento:")
                print("\n".join(sample_paras))
                print()

            # Buscar em parágrafos
            for idx, para in enumerate(doc.paragraphs):
                para_text = para.text
                if not para_text.strip():  # Pular parágrafos vazios
                    continue

                # Normalizar texto do parágrafo
                para_text_normalized = normalize_text(para_text)

                # BUSCA FLEXÍVEL: verificar se todas as palavras estão presentes (mesmo não consecutivas)
                all_words_found = all(word in para_text_normalized for word in search_words)

                # Log detalhado a cada 50 parágrafos ou quando encontrar match
                if idx % 50 == 0:
                    print(f"[DEBUG] Parágrafo {idx}: '{para_text_normalized[:80]}...'")

                # Verificar se o termo está no parágrafo
                if all_words_found:
                    print(f"[DEBUG] ✓✓✓ MATCH ENCONTRADO no parágrafo {idx}!")
                    print(f"[DEBUG] Texto original: '{para_text[:150]}'")
                    print(f"[DEBUG] Texto normalizado: '{para_text_normalized[:150]}'")

                    # Verificar formatação (negrito e sublinhado tem prioridade)
                    has_bold_underline = False

                    for run in para.runs:
                        run_text_normalized = normalize_text(run.text)
                        run_has_all_words = all(word in run_text_normalized for word in search_words)

                        if run_has_all_words:
                            if run.bold and run.underline:
                                has_bold_underline = True
                                found_contexts.append(f"[DESTAQUE] {para_text[:200]}")
                                print(f"[DEBUG] → Com formatação NEGRITO+SUBLINHADO")
                                break

                    if not has_bold_underline:
                        found_contexts.append(para_text[:200])
                        print(f"[DEBUG] → Sem formatação especial")

            # Buscar em tabelas
            print(f"[DEBUG] Total de tabelas: {len(doc.tables)}")
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text
                        if not cell_text.strip():
                            continue

                        cell_text_normalized = normalize_text(cell_text)

                        # BUSCA FLEXÍVEL em tabelas também
                        all_words_in_cell = all(word in cell_text_normalized for word in search_words)

                        if all_words_in_cell:
                            print(f"[DEBUG] ✓✓✓ MATCH em TABELA {table_idx}, linha {row_idx}, célula {cell_idx}")
                            print(f"[DEBUG] Texto: '{cell_text[:100]}'")
                            found_contexts.append(f"[TABELA] {cell_text[:200]}")

            print(f"[DEBUG] Total de contextos encontrados: {len(found_contexts)}")

            if found_contexts:
                context = " | ".join(found_contexts[:3])  # Primeiros 3 contextos
                print(f"[DEBUG] RETORNANDO SUCESSO com {len(found_contexts)} match(es)")
                return True, context

            print(f"[DEBUG] Nenhum match encontrado neste documento")

            # Se habilitado, usar IA para verificação mais profunda
            if use_ai and api_key:
                # Usar semáforo para limitar chamadas simultâneas à IA
                with self.ai_semaphore:
                    full_text = "\n".join([p.text for p in doc.paragraphs])
                    ai_result = self.check_with_ai(full_text, search_term, api_key)
                    if ai_result:
                        return True, f"[IA] Menção encontrada no contexto do documento"

            return False, ""

        except PermissionError:
            # Arquivo bloqueado
            return False, "LOCKED"
        except Exception as e:
            return False, f"ERROR: {str(e)}"

    def search_in_old_doc(self, file_path: str, search_term: str, use_ai: bool = False,
                         api_key: str = None) -> Tuple[bool, str]:
        """
        Busca em arquivos .doc antigos (formato binário)
        Usa extração de texto com 'strings'
        """
        try:
            # Extrair texto do arquivo .doc antigo
            full_text = extract_text_from_old_doc(file_path)

            if not full_text:
                print(f"[DEBUG] Não foi possível extrair texto do .doc antigo")
                return False, "ERROR: Não foi possível ler arquivo .doc antigo"

            print(f"[DEBUG] Texto extraído do .doc antigo ({len(full_text)} caracteres)")

            # Normalizar termo de busca
            search_term_normalized = normalize_text(search_term)
            search_words = search_term_normalized.split()

            print(f"[DEBUG] Buscando '{search_term}' normalizado para '{search_term_normalized}'")
            print(f"[DEBUG] Palavras: {search_words}")

            # BUSCAR NO TEXTO COMPLETO (não linha por linha, pois nomes podem estar quebrados)
            full_text_normalized = normalize_text(full_text)

            # Verificar se todas as palavras estão no documento
            all_words_in_document = all(word in full_text_normalized for word in search_words)

            if not all_words_in_document:
                print(f"[DEBUG] Nem todas as palavras encontradas no documento")
                # Mostrar quais palavras faltam
                for word in search_words:
                    if word not in full_text_normalized:
                        print(f"[DEBUG] Palavra NÃO encontrada: '{word}'")
                    else:
                        print(f"[DEBUG] Palavra encontrada: '{word}'")

            # Dividir texto em linhas e procurar contextos
            lines = full_text.split('\n')
            found_contexts = []

            # ESTRATÉGIA 1: Buscar em blocos de 3 linhas consecutivas
            # (caso o nome esteja quebrado entre linhas)
            for idx in range(len(lines)-2):
                block = ' '.join(lines[idx:idx+3])
                block_normalized = normalize_text(block)

                if all(word in block_normalized for word in search_words):
                    print(f"[DEBUG] ✓✓✓ MATCH em bloco (linhas {idx}-{idx+2})!")
                    found_contexts.append(block[:200])

            # ESTRATÉGIA 2: Buscar em linhas individuais
            for idx, line in enumerate(lines):
                if not line.strip():
                    continue

                line_normalized = normalize_text(line)

                # Busca flexível: todas as palavras devem estar presentes
                all_words_found = all(word in line_normalized for word in search_words)

                if all_words_found:
                    print(f"[DEBUG] ✓✓✓ MATCH ENCONTRADO na linha {idx}!")
                    print(f"[DEBUG] Texto: '{line[:150]}'")
                    if line[:200] not in found_contexts:  # Evitar duplicatas
                        found_contexts.append(line[:200])

            print(f"[DEBUG] Total de contextos encontrados em .doc antigo: {len(found_contexts)}")

            if found_contexts:
                context = " | ".join(found_contexts[:3])
                return True, f"[.DOC ANTIGO] {context}"

            # Se não encontrou e IA está ativada
            if use_ai and api_key:
                print(f"[DEBUG] Tentando busca com IA em .doc antigo...")
                with self.ai_semaphore:
                    ai_result = self.check_with_ai(full_text, search_term, api_key)
                    if ai_result:
                        return True, f"[IA - .DOC ANTIGO] Menção encontrada no contexto"

            return False, ""

        except Exception as e:
            print(f"[DEBUG] Erro ao processar .doc antigo: {e}")
            return False, f"ERROR: {str(e)}"

    def check_with_ai(self, text: str, search_term: str, api_key: str) -> bool:
        """
        Usa Google Gemini para verificar se o texto menciona o termo buscado
        Modelo: gemini-2.5-flash-lite
        """
        try:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-2.5-flash-lite')

            prompt = f"""
            Analise o seguinte texto e determine se há menção à pessoa ou empresa: "{search_term}"

            Considere:
            - Variações de nome (abreviações, apelidos)
            - Menções indiretas
            - Contexto de negócios/jurídico

            Responda APENAS "SIM" ou "NÃO".

            Texto:
            {text[:5000]}
            """

            response = model.generate_content(prompt)
            result = response.text.strip().upper()

            return "SIM" in result

        except Exception as e:
            print(f"Erro na verificação com IA: {e}")
            return False

    def process_files_in_thread(self, file_paths: List[str], search_term: str,
                                use_ai: bool, api_key: str, thread_id: int):
        """Processa uma lista de arquivos em uma thread"""
        print(f"[DEBUG] Thread {thread_id} iniciada com {len(file_paths)} arquivos")

        for file_path in file_paths:
            if self.stop_event.is_set():
                break

            try:
                filename = os.path.basename(file_path)
                print(f"[DEBUG] Thread {thread_id} processando: {filename}")

                found, context = self.search_in_document(file_path, search_term, use_ai, api_key)

                print(f"[DEBUG] Thread {thread_id} - {filename}: encontrado={found}, contexto={context[:50] if context else 'N/A'}")

                with self.lock:
                    self.total_files_processed += 1

                if found:
                    with self.lock:
                        self.files_found.append((file_path, context))
                        print(f"[DEBUG] Thread {thread_id} - Adicionado a files_found. Total agora: {len(self.files_found)}")
                    self.results_queue.put(('found', file_path, context))
                    print(f"[DEBUG] Thread {thread_id} - ENCONTRADO: {filename}")
                elif context == "LOCKED":
                    with self.lock:
                        self.files_with_errors.append((file_path, "Arquivo bloqueado/aberto"))
                    self.results_queue.put(('locked', file_path, None))
                elif context.startswith("ERROR:"):
                    with self.lock:
                        self.files_with_errors.append((file_path, context))
                    self.results_queue.put(('error', file_path, context))

                # Atualizar progresso
                self.progress_queue.put(('progress', self.total_files_processed))

            except Exception as e:
                print(f"[DEBUG] Thread {thread_id} ERRO: {str(e)}")
                with self.lock:
                    self.files_with_errors.append((file_path, f"Erro: {str(e)}"))
                self.results_queue.put(('error', file_path, str(e)))

        # Sinalizar que esta thread terminou
        print(f"[DEBUG] Thread {thread_id} finalizada")
        self.progress_queue.put(('thread_done', thread_id))

    def search(self, search_term: str, use_ai: bool = False, api_key: str = None) -> bool:
        """
        Inicia busca paralela
        Retorna True se iniciou com sucesso
        """
        print(f"[DEBUG] Iniciando busca...")
        print(f"[DEBUG] Termo: {search_term}")
        print(f"[DEBUG] Usar IA: {use_ai}")
        print(f"[DEBUG] Pasta base: {self.base_path}")

        self.stop_event.clear()
        self.files_found = []
        self.files_with_errors = []
        self.total_files_processed = 0

        # Verificar se o caminho existe
        if not os.path.exists(self.base_path):
            print(f"[DEBUG] ERRO: Pasta não existe: {self.base_path}")
            self.results_queue.put(('error_path', self.base_path, None))
            return False

        print(f"[DEBUG] Pasta existe, listando arquivos...")

        # Listar todos os arquivos Word recursivamente
        all_files = []
        try:
            for ext in EXTENSIONS:
                found = list(Path(self.base_path).rglob(f'*{ext}'))
                print(f"[DEBUG] Arquivos {ext}: {len(found)}")
                all_files.extend(found)
        except Exception as e:
            print(f"[DEBUG] ERRO ao listar arquivos: {e}")
            self.results_queue.put(('error_path', str(e), None))
            return False

        if not all_files:
            print(f"[DEBUG] Nenhum arquivo encontrado!")
            self.results_queue.put(('no_files', None, None))
            return False

        all_files = [str(f) for f in all_files]
        print(f"[DEBUG] Total de arquivos para processar: {len(all_files)}")

        # Dividir arquivos entre threads
        files_per_thread = len(all_files) // self.num_threads
        if files_per_thread == 0:
            files_per_thread = 1

        threads = []
        for i in range(self.num_threads):
            start_idx = i * files_per_thread
            if i == self.num_threads - 1:
                end_idx = len(all_files)
            else:
                end_idx = start_idx + files_per_thread

            thread_files = all_files[start_idx:end_idx]

            if thread_files:
                t = threading.Thread(
                    target=self.process_files_in_thread,
                    args=(thread_files, search_term, use_ai, api_key, i),
                    daemon=True
                )
                threads.append(t)
                t.start()

        # Thread de monitoramento
        def monitor_threads():
            for t in threads:
                t.join()
            print(f"[DEBUG] Todas as threads finalizadas, sinalizando conclusão...")
            self.progress_queue.put(('complete', None, None))

        monitor_thread = threading.Thread(target=monitor_threads, daemon=True)
        monitor_thread.start()

        return True

    def stop(self):
        """Para a busca"""
        self.stop_event.set()


# ===========================
# INTERFACE GRÁFICA
# ===========================
class SearchApp(ctk.CTk):
    """Aplicação principal com interface gráfica moderna"""

    def __init__(self):
        super().__init__()

        # Configurações da janela
        self.title("🔍 Busca Avançada em Transcrições - IA Integrada")
        self.geometry("1400x900")

        # Configurar tema CLARO
        ctk.set_appearance_mode("light")
        ctk.set_default_color_theme("blue")

        # Gerenciadores
        self.config_manager = ConfigManager()
        self.search_engine = WordSearchEngine(PASTA_BASE, NUM_THREADS)

        # Variáveis
        self.search_in_progress = False
        self.start_time = None
        self.progress_window = None

        # Construir interface
        self.build_ui()

        # Iniciar monitoramento de progresso
        self.after(100, self.check_progress)

    def build_ui(self):
        """Constrói a interface gráfica"""

        # ===== FRAME PRINCIPAL COM SCROLL =====
        self.main_frame = ctk.CTkScrollableFrame(self, fg_color=COLORS["bg_main"])
        self.main_frame.pack(fill="both", expand=True, padx=0, pady=0)

        # ===== CABEÇALHO =====
        header_frame = ctk.CTkFrame(self.main_frame, fg_color=COLORS["bg_header"], height=100)
        header_frame.pack(fill="x", padx=0, pady=0)
        header_frame.pack_propagate(False)

        # Container do cabeçalho (para alinhar título e botão de config)
        header_container = ctk.CTkFrame(header_frame, fg_color="transparent")
        header_container.pack(fill="both", expand=True)

        # Título centralizado
        title_label = ctk.CTkLabel(
            header_container,
            text="🔍 BUSCA AVANÇADA EM TRANSCRIÇÕES",
            font=ctk.CTkFont(size=32, weight="bold"),
            text_color=COLORS["fg_header"]
        )
        title_label.pack(pady=25)

        subtitle_label = ctk.CTkLabel(
            header_container,
            text="Powered by Google Gemini AI • 10 Threads Paralelas",
            font=ctk.CTkFont(size=12),
            text_color=COLORS["fg_header"]
        )
        subtitle_label.pack(pady=(0, 10))

        # Botão de configuração (canto superior direito)
        config_btn = ctk.CTkButton(
            header_frame,
            text="⚙️",
            width=50,
            height=50,
            font=ctk.CTkFont(size=24),
            command=self.show_config_dialog,
            fg_color=COLORS["bg_card"],
            text_color=COLORS["fg_primary"],
            hover_color=COLORS["bg_input"],
            corner_radius=25
        )
        config_btn.place(relx=0.98, rely=0.5, anchor="e")

        # ===== BUSCA =====
        search_frame = ctk.CTkFrame(self.main_frame, fg_color=COLORS["bg_card"], corner_radius=15)
        search_frame.pack(fill="x", padx=30, pady=(30, 20))

        # Container interno com padding
        search_inner = ctk.CTkFrame(search_frame, fg_color="transparent")
        search_inner.pack(fill="x", padx=20, pady=20)

        search_label = ctk.CTkLabel(
            search_inner,
            text="👤 Nome ou Empresa:",
            font=ctk.CTkFont(size=16, weight="bold"),
            text_color=COLORS["fg_primary"]
        )
        search_label.pack(anchor="w", pady=(0, 10))

        # Container para input e botão
        input_container = ctk.CTkFrame(search_inner, fg_color="transparent")
        input_container.pack(fill="x")

        self.search_entry = ctk.CTkEntry(
            input_container,
            placeholder_text="Digite o nome da pessoa ou empresa para buscar...",
            height=50,
            font=ctk.CTkFont(size=16),
            fg_color=COLORS["bg_input"],
            text_color=COLORS["fg_primary"],
            border_width=2,
            border_color=COLORS["border"],
            corner_radius=10
        )
        self.search_entry.pack(side="left", fill="x", expand=True, padx=(0, 15))
        self.search_entry.bind('<Return>', lambda e: self.start_search())

        search_btn = ctk.CTkButton(
            input_container,
            text="🔍 BUSCAR",
            width=180,
            height=50,
            command=self.start_search,
            font=ctk.CTkFont(size=16, weight="bold"),
            fg_color=COLORS["accent"],
            hover_color=COLORS["accent_hover"],
            text_color="white",
            corner_radius=10
        )
        search_btn.pack(side="right")

        # Container para checkbox de IA e info
        ai_container = ctk.CTkFrame(search_inner, fg_color="transparent")
        ai_container.pack(fill="x", pady=(15, 0))

        # Checkbox Usar IA
        self.use_ai_var = ctk.BooleanVar(value=False)
        use_ai_check = ctk.CTkCheckBox(
            ai_container,
            text="🤖 Usar IA (Google Gemini) - Mais preciso, porém mais lento",
            variable=self.use_ai_var,
            font=ctk.CTkFont(size=13, weight="bold"),
            text_color=COLORS["fg_primary"],
            fg_color=COLORS["accent"],
            hover_color=COLORS["accent_hover"]
        )
        use_ai_check.pack(side="left")

        # Info sobre formatação
        format_info_label = ctk.CTkLabel(
            search_inner,
            text="💡 Prioriza termos em Negrito e Sublinhado • Ignora acentos automaticamente",
            font=ctk.CTkFont(size=11),
            text_color=COLORS["fg_secondary"]
        )
        format_info_label.pack(anchor="w", pady=(5, 0))

        # ===== INFO PASTA =====
        info_frame = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        info_frame.pack(fill="x", padx=30, pady=(0, 10))

        info_label = ctk.CTkLabel(
            info_frame,
            text=f"📁 Pasta Base: {PASTA_BASE}",
            font=ctk.CTkFont(size=11),
            text_color=COLORS["fg_secondary"]
        )
        info_label.pack(anchor="w")

        # ===== CONTAINER DE RESULTADOS =====
        results_container = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        results_container.pack(fill="both", expand=True, padx=30, pady=(10, 20))

        # Configurar grid
        results_container.grid_columnconfigure(0, weight=1)
        results_container.grid_columnconfigure(1, weight=1)
        results_container.grid_rowconfigure(0, weight=1)

        # ===== PAINEL DE RESULTADOS =====
        results_frame = ctk.CTkFrame(results_container, fg_color=COLORS["bg_card"], corner_radius=15)
        results_frame.grid(row=0, column=0, sticky="nsew", padx=(0, 10))

        results_header = ctk.CTkFrame(results_frame, fg_color=COLORS["success"], corner_radius=10, height=50)
        results_header.pack(fill="x", padx=15, pady=15)
        results_header.pack_propagate(False)

        results_title = ctk.CTkLabel(
            results_header,
            text="📋 RESULTADOS",
            font=ctk.CTkFont(size=18, weight="bold"),
            text_color="white"
        )
        results_title.pack(pady=10)

        # Frame scrollable para resultados modernos
        self.results_scrollable = ctk.CTkScrollableFrame(
            results_frame,
            fg_color=COLORS["bg_input"],
            corner_radius=10
        )
        self.results_scrollable.pack(fill="both", expand=True, padx=15, pady=(0, 15))

        # ===== PAINEL DE ERROS =====
        errors_frame = ctk.CTkFrame(results_container, fg_color=COLORS["bg_card"], corner_radius=15)
        errors_frame.grid(row=0, column=1, sticky="nsew", padx=(10, 0))

        errors_header = ctk.CTkFrame(errors_frame, fg_color=COLORS["warning"], corner_radius=10, height=50)
        errors_header.pack(fill="x", padx=15, pady=15)
        errors_header.pack_propagate(False)

        errors_title = ctk.CTkLabel(
            errors_header,
            text="⚠️ ARQUIVOS NÃO ACESSADOS",
            font=ctk.CTkFont(size=18, weight="bold"),
            text_color="white"
        )
        errors_title.pack(pady=10)

        self.errors_textbox = ctk.CTkTextbox(
            errors_frame,
            font=ctk.CTkFont(size=12),
            fg_color=COLORS["bg_input"],
            text_color=COLORS["fg_secondary"],
            wrap="word",
            corner_radius=10
        )
        self.errors_textbox.pack(fill="both", expand=True, padx=15, pady=(0, 15))

        # ===== STATUS BAR =====
        status_frame = ctk.CTkFrame(self.main_frame, fg_color=COLORS["bg_header"], height=50)
        status_frame.pack(fill="x", padx=0, pady=0)
        status_frame.pack_propagate(False)

        self.status_label = ctk.CTkLabel(
            status_frame,
            text="✅ Sistema pronto para busca...",
            font=ctk.CTkFont(size=13),
            text_color=COLORS["fg_header"]
        )
        self.status_label.pack(side="left", padx=30, pady=15)

    def show_config_dialog(self):
        """Mostra janela de configuração da API Key"""
        # Criar janela modal
        config_window = ctk.CTkToplevel(self)
        config_window.title("⚙️ Configurações")
        config_window.geometry("650x450")
        config_window.resizable(False, False)

        # Centralizar janela em relação à janela principal
        config_window.transient(self)
        config_window.grab_set()

        # Atualizar para obter dimensões corretas
        config_window.update_idletasks()

        # Centralizar
        x = self.winfo_x() + (self.winfo_width() // 2) - (650 // 2)
        y = self.winfo_y() + (self.winfo_height() // 2) - (450 // 2)
        config_window.geometry(f"650x450+{x}+{y}")

        # Frame principal
        main_config_frame = ctk.CTkFrame(config_window, fg_color=COLORS["bg_main"])
        main_config_frame.pack(fill="both", expand=True, padx=0, pady=0)

        # Título
        title_config = ctk.CTkLabel(
            main_config_frame,
            text="🔑 Configuração da API Key",
            font=ctk.CTkFont(size=26, weight="bold"),
            text_color=COLORS["fg_primary"]
        )
        title_config.pack(pady=30)

        # Frame do conteúdo
        content_frame = ctk.CTkFrame(main_config_frame, fg_color=COLORS["bg_card"], corner_radius=15)
        content_frame.pack(fill="both", expand=True, padx=40, pady=(10, 30))

        # Instrução
        instruction_label = ctk.CTkLabel(
            content_frame,
            text="Insira sua API Key do Google Gemini:",
            font=ctk.CTkFont(size=16, weight="bold"),
            text_color=COLORS["fg_primary"]
        )
        instruction_label.pack(pady=(30, 15), padx=30)

        # Campo de entrada
        api_key_entry = ctk.CTkEntry(
            content_frame,
            placeholder_text="Cole sua API Key aqui...",
            width=520,
            height=50,
            font=ctk.CTkFont(size=14),
            fg_color=COLORS["bg_input"],
            text_color=COLORS["fg_primary"],
            border_width=2,
            border_color=COLORS["border"],
            corner_radius=10
        )
        api_key_entry.pack(pady=15, padx=30)

        # Carregar API Key salva
        saved_key = self.config_manager.get_api_key()
        if saved_key:
            api_key_entry.insert(0, saved_key)

        # Link para obter API Key
        link_label = ctk.CTkLabel(
            content_frame,
            text="🔗 Obter API Key: https://makersuite.google.com/app/apikey",
            font=ctk.CTkFont(size=13),
            text_color=COLORS["info"],
            cursor="hand2"
        )
        link_label.pack(pady=(10, 20), padx=30)

        # Função para salvar
        def save_and_close():
            api_key = api_key_entry.get().strip()
            if not api_key:
                messagebox.showwarning("Atenção", "Por favor, insira a API Key.")
                return

            if self.config_manager.set_api_key(api_key):
                messagebox.showinfo("Sucesso", "API Key salva com sucesso!")
                self.status_label.configure(text="✅ API Key configurada e salva")
                config_window.destroy()
            else:
                messagebox.showerror("Erro", "Erro ao salvar API Key")

        # Botões
        buttons_frame = ctk.CTkFrame(content_frame, fg_color="transparent")
        buttons_frame.pack(pady=(15, 30))

        save_btn = ctk.CTkButton(
            buttons_frame,
            text="💾 Salvar",
            width=200,
            height=50,
            command=save_and_close,
            font=ctk.CTkFont(size=16, weight="bold"),
            fg_color=COLORS["success"],
            hover_color="#229954",
            text_color="white",
            corner_radius=10
        )
        save_btn.pack(side="left", padx=15)

        cancel_btn = ctk.CTkButton(
            buttons_frame,
            text="✕ Cancelar",
            width=200,
            height=50,
            command=config_window.destroy,
            font=ctk.CTkFont(size=16),
            fg_color=COLORS["error"],
            hover_color="#C0392B",
            text_color="white",
            corner_radius=10
        )
        cancel_btn.pack(side="left", padx=15)

    def start_search(self):
        """Inicia o processo de busca"""
        if self.search_in_progress:
            messagebox.showwarning("Atenção", "Uma busca já está em andamento!")
            return

        search_term = self.search_entry.get().strip()
        if not search_term:
            messagebox.showwarning("Atenção", "Por favor, digite um nome ou empresa para buscar.")
            return

        # Verificar se usuário quer usar IA
        use_ai = self.use_ai_var.get()
        api_key = self.config_manager.get_api_key()

        # Se IA ativada, verificar se tem API Key
        if use_ai and not api_key:
            messagebox.showwarning(
                "Configuração Necessária",
                "Para usar IA, configure a API Key do Gemini clicando no botão ⚙️ no canto superior direito.\n\n"
                "Ou desmarque a opção 'Usar IA' para busca apenas textual."
            )
            return

        # Limpar resultados anteriores
        self._clear_results()
        self.errors_textbox.delete("1.0", "end")

        # Armazenar termo de busca para usar em finish_search
        self.current_search_term = search_term

        # Atualizar status
        self.search_in_progress = True
        self.start_time = time.time()
        if use_ai:
            self.status_label.configure(text="🔄 Buscando arquivos com IA...")
        else:
            self.status_label.configure(text="🔄 Buscando arquivos (modo rápido)...")

        # Mostrar janela de progresso
        self.show_progress_window()

        # CRÍTICO: Reiniciar o check_progress loop
        print(f"[DEBUG] Agendando check_progress para monitorar busca...")
        self.after(100, self.check_progress)

        # Iniciar busca em thread separada
        search_thread = threading.Thread(
            target=self.search_engine.search,
            args=(search_term, use_ai, api_key),
            daemon=True
        )
        search_thread.start()

    def show_progress_window(self):
        """Mostra janela de progresso modal"""
        self.progress_window = ctk.CTkToplevel(self)
        self.progress_window.title("Processando...")
        self.progress_window.geometry("450x300")
        self.progress_window.resizable(False, False)

        # Centralizar janela
        self.progress_window.transient(self)
        self.progress_window.grab_set()

        # Frame principal
        progress_main = ctk.CTkFrame(self.progress_window, fg_color=COLORS["bg_card"])
        progress_main.pack(fill="both", expand=True)

        # Ícone animado
        self.progress_icon = ctk.CTkLabel(
            progress_main,
            text="⏳",
            font=ctk.CTkFont(size=60)
        )
        self.progress_icon.pack(pady=(30, 10))

        # Mensagem (varia conforme uso de IA)
        use_ai = self.use_ai_var.get()
        message_text = "Processando arquivos com IA..." if use_ai else "Processando arquivos..."

        self.progress_message = ctk.CTkLabel(
            progress_main,
            text=message_text,
            font=ctk.CTkFont(size=18, weight="bold"),
            text_color=COLORS["fg_primary"]
        )
        self.progress_message.pack(pady=10)

        # Contador de arquivos
        self.progress_count = ctk.CTkLabel(
            progress_main,
            text="Arquivos encontrados: 0",
            font=ctk.CTkFont(size=15),
            text_color=COLORS["fg_secondary"]
        )
        self.progress_count.pack(pady=5)

        # Tempo decorrido
        self.progress_time = ctk.CTkLabel(
            progress_main,
            text="Tempo: 00:00",
            font=ctk.CTkFont(size=15),
            text_color=COLORS["fg_secondary"]
        )
        self.progress_time.pack(pady=5)

        # Barra de progresso indeterminada
        self.progress_bar = ctk.CTkProgressBar(
            progress_main,
            mode="indeterminate",
            width=380,
            height=10,
            progress_color=COLORS["accent"]
        )
        self.progress_bar.pack(pady=(20, 30))
        self.progress_bar.start()

    def check_progress(self):
        """Verifica o progresso da busca periodicamente"""
        # Log de execução periódica (a cada 10 chamadas para não poluir)
        if not hasattr(self, '_check_progress_count'):
            self._check_progress_count = 0
        self._check_progress_count += 1
        if self._check_progress_count % 10 == 0:
            print(f"[DEBUG] check_progress rodando... (chamada #{self._check_progress_count}, search_in_progress={self.search_in_progress})")

        try:
            while True:
                msg_type, data, extra = self.search_engine.progress_queue.get_nowait()
                print(f"[DEBUG] check_progress recebeu: msg_type={msg_type}, data={data}")

                if msg_type == 'complete':
                    print(f"[DEBUG] Mensagem 'complete' recebida, chamando finish_search()...")
                    self.finish_search()
                    break
                elif msg_type == 'progress':
                    # Atualizar contadores no progress window
                    if self.progress_window and self.progress_window.winfo_exists():
                        elapsed = time.time() - self.start_time
                        minutes = int(elapsed // 60)
                        seconds = int(elapsed % 60)

                        self.progress_count.configure(
                            text=f"Arquivos encontrados: {len(self.search_engine.files_found)}"
                        )
                        self.progress_time.configure(
                            text=f"Tempo: {minutes:02d}:{seconds:02d}"
                        )
        except queue.Empty:
            pass
        except Exception as e:
            print(f"[DEBUG] ERRO em check_progress: {e}")
            import traceback
            traceback.print_exc()

        # Continuar verificando
        if self.search_in_progress:
            self.after(100, self.check_progress)
        else:
            print(f"[DEBUG] check_progress parando porque search_in_progress=False")

    def _extract_name_from_context(self, context: str, search_term: str) -> str:
        """Extrai apenas o nome encontrado do contexto"""
        if not context or context == "N/A":
            return search_term.upper()

        # Remover prefixos técnicos
        context = context.replace("[.DOC ANTIGO] ", "")
        context = context.replace("[IA - .DOC ANTIGO] ", "")
        context = context.replace("Menção encontrada no contexto", "")

        # Limpar e normalizar
        context = ' '.join(context.split())

        # Procurar por sequências de palavras em maiúsculas (geralmente nomes)
        import re
        # Padrão: sequência de 2-5 palavras em maiúsculas
        pattern = r'\b[A-ZÁÉÍÓÚÂÊÔÃÕÇ][A-ZÁÉÍÓÚÂÊÔÃÕÇ\s]{2,50}\b'
        matches = re.findall(pattern, context)

        if matches:
            # Retornar a primeira sequência encontrada que contenha pelo menos uma palavra do termo de busca
            search_words = normalize_text(search_term).split()
            for match in matches:
                match_normalized = normalize_text(match)
                if any(word in match_normalized for word in search_words):
                    return match.strip()

        # Se não encontrou, retornar o termo de busca em maiúsculas
        return search_term.upper()

    def _clear_results(self):
        """Limpa os resultados anteriores"""
        # Destruir todos os widgets do frame scrollable
        for widget in self.results_scrollable.winfo_children():
            widget.destroy()

    def finish_search(self):
        """Finaliza a busca e exibe resultados"""
        print(f"[DEBUG] finish_search() chamada!")
        print(f"[DEBUG] Arquivos encontrados: {len(self.search_engine.files_found)}")
        print(f"[DEBUG] Arquivos com erro: {len(self.search_engine.files_with_errors)}")

        self.search_in_progress = False

        # Fechar janela de progresso
        if self.progress_window and self.progress_window.winfo_exists():
            self.progress_bar.stop()
            self.progress_window.destroy()

        # Calcular tempo total
        elapsed = time.time() - self.start_time
        minutes = int(elapsed // 60)
        seconds = int(elapsed % 60)

        # Exibir resultados
        num_found = len(self.search_engine.files_found)
        num_errors = len(self.search_engine.files_with_errors)

        print(f"[DEBUG] Exibindo {num_found} resultados na GUI...")

        # Limpar resultados anteriores
        self._clear_results()

        if num_found > 0:
            # Header com contador
            header_label = ctk.CTkLabel(
                self.results_scrollable,
                text=f"✅ {num_found} arquivo(s) encontrado(s)!",
                font=ctk.CTkFont(size=16, weight="bold"),
                text_color=COLORS["success"]
            )
            header_label.pack(pady=(10, 20), anchor="w")

            # Criar um botão moderno para cada resultado
            for i, (file_path, context) in enumerate(self.search_engine.files_found, 1):
                filename = os.path.basename(file_path)

                # Extrair apenas o nome encontrado
                name_found = self._extract_name_from_context(context, self.current_search_term)

                # Frame para cada resultado
                result_frame = ctk.CTkFrame(
                    self.results_scrollable,
                    fg_color=COLORS["bg_card"],
                    corner_radius=10,
                    height=60
                )
                result_frame.pack(fill="x", pady=5)
                result_frame.pack_propagate(False)

                # Botão do arquivo (ocupa ~50% do espaço)
                file_button = ctk.CTkButton(
                    result_frame,
                    text=f"📄 {filename}",
                    font=ctk.CTkFont(size=13, weight="bold"),
                    fg_color=COLORS["accent"],
                    hover_color=COLORS["accent_hover"],
                    text_color="white",
                    corner_radius=8,
                    command=lambda fp=file_path: self.open_file(fp),
                    anchor="w"
                )
                file_button.pack(side="left", fill="both", expand=True, padx=10, pady=10)

                # Label com o nome encontrado (ocupa ~50% do espaço)
                name_label = ctk.CTkLabel(
                    result_frame,
                    text=f"→  {name_found}",
                    font=ctk.CTkFont(size=13),
                    text_color=COLORS["fg_primary"],
                    anchor="w"
                )
                name_label.pack(side="left", fill="both", expand=True, padx=(0, 10), pady=10)

            # Dica no final
            tip_label = ctk.CTkLabel(
                self.results_scrollable,
                text="💡 Clique no botão do arquivo para abrir",
                font=ctk.CTkFont(size=11),
                text_color=COLORS["info"]
            )
            tip_label.pack(pady=(20, 10), anchor="w")

            self.status_label.configure(
                text=f"✅ Busca concluída! {num_found} arquivo(s) encontrado(s) em {minutes:02d}:{seconds:02d}"
            )
        else:
            # Mensagem de nenhum resultado
            no_result_label = ctk.CTkLabel(
                self.results_scrollable,
                text="❌ Nenhum arquivo encontrado",
                font=ctk.CTkFont(size=16, weight="bold"),
                text_color=COLORS["error"]
            )
            no_result_label.pack(pady=50)
            self.status_label.configure(text=f"⚠️ Nenhum resultado encontrado em {minutes:02d}:{seconds:02d}")

        # Exibir erros
        if num_errors > 0:
            self.errors_textbox.insert("end", f"⚠️ {num_errors} arquivo(s) não acessado(s):\n\n")

            for i, (file_path, error_msg) in enumerate(self.search_engine.files_with_errors, 1):
                filename = os.path.basename(file_path)
                self.errors_textbox.insert("end", f"{i}. 🔒 {filename}\n")
                self.errors_textbox.insert("end", f"   Motivo: {error_msg}\n\n")
        else:
            self.errors_textbox.insert("end", "✅ Todos os arquivos foram acessados com sucesso!")

        # Mostrar popup de conclusão
        messagebox.showinfo(
            "Busca Concluída",
            f"Busca finalizada em {minutes:02d}:{seconds:02d}\n\n"
            f"✅ Encontrados: {num_found}\n"
            f"⚠️ Erros: {num_errors}"
        )

    def open_file_from_selection(self, event):
        """Abre arquivo ao clicar duas vezes"""
        try:
            # Pegar a linha clicada
            index = self.results_textbox.index("@%s,%s" % (event.x, event.y))
            line_start = index.split('.')[0]
            line_content = self.results_textbox.get(f"{line_start}.0", f"{line_start}.end")

            # Buscar o arquivo correspondente
            for file_path, _ in self.search_engine.files_found:
                if file_path in line_content or os.path.basename(file_path) in line_content:
                    self.open_file(file_path)
                    break
        except Exception as e:
            print(f"Erro ao abrir arquivo: {e}")

    def open_file(self, file_path: str):
        """Abre um arquivo Word"""
        try:
            if os.path.exists(file_path):
                if sys.platform == "win32":
                    os.startfile(file_path)
                elif sys.platform == "darwin":  # macOS
                    os.system(f'open "{file_path}"')
                else:  # Linux
                    os.system(f'xdg-open "{file_path}"')

                self.status_label.configure(text=f"📄 Arquivo aberto: {os.path.basename(file_path)}")
            else:
                messagebox.showerror("Erro", "Arquivo não encontrado!")
        except Exception as e:
            messagebox.showerror("Erro", f"Não foi possível abrir o arquivo:\n{str(e)}")


# ===========================
# FUNÇÃO PRINCIPAL
# ===========================
def main():
    """Função principal"""
    print("=" * 60)
    print("🔍 Sistema de Busca Avançada em Arquivos Word com IA")
    print("=" * 60)
    print(f"📁 Pasta Base: {PASTA_BASE}")
    print(f"🧵 Threads: {NUM_THREADS}")
    print(f"📝 Extensões: {', '.join(EXTENSIONS)}")
    print(f"🤖 IA: Opcional (Google Gemini)")
    print("=" * 60)
    print()

    # Teste de normalização
    print("[TESTE] Testando função de normalização:")
    test_cases = [
        "José Silva",
        "ANTONIO CARDINÁ",
        "João Paulo",
        "Café",
        "São Paulo"
    ]
    for test in test_cases:
        normalized = normalize_text(test)
        print(f"  '{test}' → '{normalized}'")
    print()

    # Verificar se a pasta existe
    if not os.path.exists(PASTA_BASE):
        print(f"⚠️ AVISO: Pasta base não encontrada: {PASTA_BASE}")
        print("O programa continuará, mas você pode precisar ajustar o caminho.")
        print()

    # Iniciar aplicação
    app = SearchApp()
    app.mainloop()


if __name__ == "__main__":
    main()
